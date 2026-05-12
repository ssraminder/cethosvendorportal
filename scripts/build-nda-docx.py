"""
Build .docx versions of the NDA templates from the HTML stored in
`nda_templates` for offline legal/HR reference. Run with:

    python scripts/build-nda-docx.py

Writes one file per active template into docs/nda-templates/.

The HTML schema is intentionally narrow (h2/h3/p/strong) so a minimal
HTML-to-Word pass is enough — no need to pull in a full converter.
"""
from __future__ import annotations

import re
import sys
from datetime import datetime, timezone
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Hard-coded so we don't need to hit the DB at build time. Bump these
# verbatim from `SELECT version_label, body_html FROM nda_templates`
# whenever a new version goes live.
TEMPLATES = [
    {
        "version_label": "v1.0",
        "title": "Cethos Translator Confidentiality & Non-Disclosure Agreement",
        "effective_from": "2026-05-12",
        "body_html": """<h2>Cethos Translator Confidentiality &amp; Non-Disclosure Agreement</h2>
<p>This Confidentiality and Non-Disclosure Agreement ("Agreement") is entered into between <strong>Cethos Translation Services</strong> ("Cethos") and the undersigned vendor ("Translator") as of the date of acceptance below.</p>

<h3>1. Confidential Information</h3>
<p>"Confidential Information" means any non-public information disclosed by Cethos or its clients to the Translator in connection with translation, revision, review, proofreading, transcription, interpretation, or related services. This includes source documents, glossaries, translation memories, style guides, client identities, project metadata, pricing, and any other business or technical information not generally known to the public.</p>

<h3>2. Obligations</h3>
<p>The Translator agrees to: (a) keep all Confidential Information strictly confidential; (b) use it solely to perform the agreed services for Cethos; (c) not disclose it to any third party without Cethos's prior written consent, including any subcontractor, family member, friend, or AI/machine-translation service that retains or trains on input; (d) protect it with at least the same care used to protect their own confidential information, and never less than a reasonable standard of care; (e) destroy or return all copies upon Cethos's request or upon completion of services.</p>

<h3>3. AI and machine-translation tools</h3>
<p>Translator shall not input Confidential Information into any public or third-party AI or machine-translation service (including but not limited to Google Translate, DeepL Free, ChatGPT, Claude, Gemini) unless explicitly authorized in writing by Cethos for the specific project. Approved CAT tools with private, non-training infrastructure (e.g., Trados, MemoQ, Phrase) may be used per project instructions.</p>

<h3>4. Personal data and GDPR</h3>
<p>Where Confidential Information includes personal data (names, contact details, medical information, identity documents, etc.), Translator acts as a processor on behalf of Cethos and must comply with applicable data-protection law (including GDPR where the data subject is in the EU/UK). Translator shall promptly notify Cethos of any actual or suspected personal-data breach.</p>

<h3>5. Term and survival</h3>
<p>This Agreement takes effect on the date of acceptance and continues for so long as Translator engages with Cethos. The obligations of confidentiality survive termination indefinitely for trade secrets and for a period of five (5) years for all other Confidential Information.</p>

<h3>6. No grant of rights</h3>
<p>Nothing in this Agreement grants Translator any license, ownership, or other right in the Confidential Information except the limited right to use it to perform the agreed services.</p>

<h3>7. Remedies</h3>
<p>Translator acknowledges that any breach may cause irreparable harm to Cethos and its clients, and that monetary damages alone may be insufficient. Cethos may seek injunctive relief in addition to any other remedies available at law or in equity.</p>

<h3>8. Governing law</h3>
<p>This Agreement is governed by the laws of the Province of Alberta, Canada, without regard to its conflict-of-laws principles. Disputes shall be resolved in the courts of Alberta.</p>

<h3>9. Acceptance</h3>
<p>By typing your full legal name below and clicking "I agree", you confirm that you have read, understood, and agreed to be bound by this Agreement. You acknowledge that this electronic acceptance has the same legal effect as a handwritten signature.</p>""",
    },
    {
        "version_label": "v1.1",
        "title": "Cethos Translator Confidentiality & Non-Disclosure Agreement",
        "effective_from": "2026-05-12",
        "body_html": """<h2>Cethos Translator Confidentiality &amp; Non-Disclosure Agreement</h2>
<p>This Confidentiality and Non-Disclosure Agreement ("Agreement") is entered into between <strong>Cethos Translation Services</strong> ("Cethos") and the undersigned vendor ("Translator") as of the date of acceptance below.</p>

<h3>1. Confidential Information</h3>
<p>"Confidential Information" means any non-public information disclosed by Cethos or its clients to the Translator in connection with translation, revision, review, proofreading, transcription, interpretation, or related services. This includes source documents, glossaries, translation memories, style guides, client identities, project metadata, pricing, and any other business or technical information not generally known to the public.</p>

<h3>2. Obligations</h3>
<p>The Translator agrees to: (a) keep all Confidential Information strictly confidential; (b) use it solely to perform the agreed services for Cethos; (c) not disclose it to any third party without Cethos's prior written consent, including any subcontractor, family member, friend, or AI/machine-translation service that retains or trains on input; (d) protect it with at least the same care used to protect their own confidential information, and never less than a reasonable standard of care; (e) destroy or return all copies upon Cethos's request or upon completion of services.</p>

<h3>3. AI and machine-translation tools</h3>
<p>Translator shall not input Confidential Information into any public or third-party AI or machine-translation service (including but not limited to Google Translate, DeepL Free, ChatGPT, Claude, Gemini) unless explicitly authorized in writing by Cethos for the specific project. Approved CAT tools with private, non-training infrastructure (e.g., Trados, MemoQ, Phrase) may be used per project instructions.</p>

<h3>4. Personal data and GDPR</h3>
<p>Where Confidential Information includes personal data (names, contact details, medical information, identity documents, etc.), Translator acts as a processor on behalf of Cethos and must comply with applicable data-protection law (including GDPR where the data subject is in the EU/UK). Translator shall promptly notify Cethos of any actual or suspected personal-data breach.</p>

<h3>5. Term and survival</h3>
<p>This Agreement takes effect on the date of acceptance and continues for so long as Translator engages with Cethos. The obligations of confidentiality survive termination indefinitely for trade secrets and for a period of five (5) years for all other Confidential Information.</p>

<h3>6. No grant of rights</h3>
<p>Nothing in this Agreement grants Translator any license, ownership, or other right in the Confidential Information except the limited right to use it to perform the agreed services.</p>

<h3>7. Remedies</h3>
<p>Translator acknowledges that any breach may cause irreparable harm to Cethos and its clients, and that monetary damages alone may be insufficient. Cethos may seek injunctive relief in addition to any other remedies available at law or in equity.</p>

<h3>8. Governing law</h3>
<p>This Agreement is governed by the laws of the Province of Alberta, Canada, without regard to its conflict-of-laws principles. Disputes shall be resolved in the courts of Alberta.</p>

<h3>9. Acceptance</h3>
<p>By typing your full legal name below and clicking "I agree", you confirm that you have read, understood, and agreed to be bound by this Agreement. You acknowledge that this electronic acceptance has the same legal effect as a handwritten signature.</p>""",
    },
    {
        "version_label": "v2.3",
        "title": "Cethos Contractor Confidentiality & Non-Disclosure Agreement",
        "effective_from": "2026-05-12",
        "body_html": """<h2>Cethos Contractor Confidentiality &amp; Non-Disclosure Agreement</h2>
<p>This Confidentiality and Non-Disclosure Agreement (“Agreement”) is entered into between Cethos and the undersigned contractor, vendor, company, or individual service provider (“Contractor”) as of the date of acceptance below.</p>
<p>For purposes of this Agreement, “Cethos” means Cethos Solutions Inc., 12537494 Canada Inc., and each of their present and future Affiliates and Subsidiaries.</p>
<p>“Affiliate” means, with respect to an entity, any other entity that directly or indirectly controls, is controlled by, or is under common control with that entity, whether now existing or hereafter created or acquired.</p>
<p>“Subsidiary” means any entity that is directly or indirectly controlled by another entity.</p>
<p>“Control” means the direct or indirect power to direct or cause the direction of the management, policies, or affairs of an entity, whether through ownership of voting securities, by contract, or otherwise.</p>
<h3>1. Purpose and scope</h3>
<p>Contractor may receive access to Confidential Information in connection with services performed for Cethos or for Cethos clients, whether directly or through a subcontracting chain. This Agreement applies to any Contractor performing language or related services, including translators, reviewers, editors, proofreaders, transcribers, interpreters, DTP artists, localization engineers, project support personnel, and any other personnel engaged by Contractor to perform such services.</p>
<h3>2. Confidential Information</h3>
<p>“Confidential Information” means any non-public information disclosed or made available by Cethos, its clients, or any upstream contractor in connection with the services. Confidential Information includes source documents, target content, glossaries, terminology databases, translation memories, style guides, client identities, project metadata, workflows, instructions, pricing, technical, commercial, legal, medical, financial, or personal data, and any other information that a reasonable person would understand to be confidential.</p>
<p>Confidential Information does not include information that Contractor can demonstrate by written records: (a) was already lawfully known to Contractor without restriction before disclosure; (b) becomes public through no breach of this Agreement; (c) is lawfully received from a third party without a duty of confidentiality; or (d) is independently developed without use of or reference to the Confidential Information.</p>
<h3>3. Confidentiality obligations</h3>
<p>Contractor shall:</p>
<ul>
  <li>keep all Confidential Information strictly confidential;</li>
  <li>use Confidential Information solely to perform the agreed services;</li>
  <li>not disclose Confidential Information to any third party except to Contractor personnel or approved subcontractors who have a strict need to know for the services and who are bound by written confidentiality obligations no less protective than this Agreement;</li>
  <li>protect Confidential Information with at least the same degree of care used to protect Contractor’s own confidential information of a similar nature, and in no event less than a reasonable standard of care;</li>
  <li>notify Cethos without undue delay, and in any event within twenty-four (24) hours of becoming aware, of any unauthorized access, use, or disclosure of Confidential Information, with such notice given in accordance with Section 18 (Notices); and</li>
  <li>upon request from Cethos or upon completion of the services, and in any event within thirty (30) days, return or securely delete all Confidential Information in Contractor’s possession or control, and certify such return or deletion in writing on Cethos’s request. The foregoing does not apply to archival backup copies maintained automatically in the ordinary course of business and not readily accessible for active use, provided such copies remain protected under this Agreement until overwritten or deleted in the normal backup cycle.</li>
</ul>
<h3>4. Subcontracting and downstream personnel</h3>
<p>If Contractor is registered in Cethos’s vendor records as an approved business entity, Contractor may use its own employees, in-house linguists, affiliates, freelance subcontractors, or other downstream personnel to perform the services, but only to the extent reasonably necessary and only if such persons are bound by written confidentiality and data-protection obligations no less protective than this Agreement.</p>
<p>If Contractor is registered in Cethos’s vendor records as an individual, Contractor shall perform the services personally and shall not subcontract, delegate, or transfer any part of the services to any third party without Cethos’s prior written approval for the relevant project.</p>
<p>Contractor remains fully responsible for any act or omission of its employees, in-house personnel, affiliates, subcontractors, and downstream freelancers relating to the services or to Confidential Information, to the same extent as if the act or omission were Contractor’s own.</p>
<p>For purposes of this Section, the Contractor’s status as an approved business entity or an individual is determined by the vendor profile completed by Contractor during onboarding and maintained by Cethos. On reasonable request, Cethos will confirm a Contractor’s then-current status in writing.</p>
<p>Cethos may also specify in project instructions that prior approval is required for any further subcontracting on a particular project, and Contractor shall comply with those instructions.</p>
<h3>5. AI and machine-translation tools</h3>
<p>Contractor shall not input Confidential Information into any public or third-party AI, generative AI, or machine-translation service that retains, trains on, or otherwise uses submitted content beyond providing the requested service, unless Cethos has expressly authorized such use in writing for the relevant project or category of projects.</p>
<p>Contractor may use CAT tools, secure machine-translation systems, terminology platforms, QA tools, and similar technology where: (a) the tool is operated under commercially reasonable confidentiality and security protections; (b) submitted content is not used to train public models or otherwise made available to unauthorized parties; and (c) such use is not prohibited by Cethos project instructions, client requirements, or Cethos’s published tool and security guidelines as made available to Contractor from time to time.</p>
<p>Contractor is responsible for ensuring that any technology used by it or its downstream personnel complies with this section.</p>
<p>By way of illustration and without limitation: Contractor shall not input Confidential Information into consumer-tier or free public services that retain or train on user inputs (such as the consumer tier of ChatGPT, Google Translate’s free public web interface, or DeepL Free), unless expressly authorized by Cethos in writing for the relevant project. Permitted categories typically include CAT tools (such as memoQ, Trados Studio, or Phrase) operating locally or under enterprise licences, enterprise machine-translation services configured with no-train or zero-retention settings, and Cethos-provisioned environments. Specific permitted and prohibited services may be set out in Cethos’s published tool guidelines as updated from time to time.</p>
<h3>6. Personal data and security</h3>
<p>To the extent Confidential Information includes personal data, Contractor shall process such personal data only for the purpose of performing the services and in accordance with applicable data-protection law (including, where applicable, the Personal Information Protection and Electronic Documents Act (Canada), the Personal Information Protection Act (Alberta), the Health Information Act (Alberta), and the EU and UK General Data Protection Regulation), and in accordance with Cethos instructions. Contractor shall implement appropriate technical and organizational measures to protect personal data and shall notify Cethos without undue delay of any actual or suspected personal-data breach affecting Confidential Information.</p>
<p>Where required by applicable law or by client instructions, the parties may enter into a separate data-processing agreement that supplements this Agreement.</p>
<h3>7. Compelled disclosure</h3>
<p>If Contractor is legally required by law, regulation, court order, or governmental demand to disclose Confidential Information, Contractor shall, to the extent legally permitted, promptly notify Cethos before disclosure so that Cethos may seek a protective order or other appropriate remedy. Contractor shall disclose only the minimum portion of Confidential Information legally required.</p>
<h3>8. Ownership and limited use</h3>
<p>All Confidential Information remains the property of Cethos, its clients, or the applicable owner. Nothing in this Agreement grants Contractor any license, ownership interest, or other rights in Confidential Information except the limited right to use it as necessary to perform the services.</p>
<p>Ownership of intellectual property in any work product produced by Contractor in the course of the services — including translated, edited, reviewed, transcribed, interpreted, or localized content — is governed by the applicable services agreement, purchase order, or project terms between the parties. Where no such agreement specifies otherwise, Contractor hereby assigns to Cethos, on creation, all right, title, and interest, including copyright, in any such work product, and waives in favour of Cethos all moral rights therein to the extent permitted by law.</p>
<h3>9. Term and survival</h3>
<p>This Agreement takes effect on the date of acceptance and continues for so long as Contractor performs services for Cethos or receives Confidential Information in connection with potential or actual services.</p>
<p>The obligations of confidentiality and restricted use survive termination of the parties’ relationship and continue: (a) indefinitely for trade secrets — meaning information that derives independent economic value, actual or potential, from not being generally known to, and not being readily ascertainable by proper means by, other persons who could obtain economic value from its disclosure or use, and that is the subject of reasonable efforts by Cethos or its clients to maintain its secrecy — for so long as such information remains a trade secret under applicable law; and (b) for five (5) years after disclosure for all other Confidential Information, unless a longer period is required by applicable law, client requirements, or a project-specific agreement.</p>
<h3>10. Non-solicitation and non-circumvention</h3>
<p>During the term of this Agreement and for a period of twelve (12) months following the end of Contractor’s relationship with Cethos, Contractor shall not, directly or indirectly, on its own behalf or on behalf of any other person or entity:</p>
<ul>
  <li>solicit, divert, or attempt to solicit or divert from Cethos any client or active prospective client of Cethos with whom Contractor had direct contact or about whom Contractor received Confidential Information in connection with the services;</li>
  <li>solicit, induce, recruit, or encourage any employee, in-house linguist, freelancer, subcontractor, or other contractor of Cethos with whom Contractor had material dealings in connection with the services to reduce or terminate their relationship with Cethos; or</li>
  <li>knowingly assist any other person or entity to do the foregoing.</li>
</ul>
<p>For purposes of this Section, “active prospective client” means any person or entity that, within the twelve (12) months preceding the relevant date, (a) received a quote, proposal, or pricing from Cethos; (b) is recorded as an open opportunity in Cethos’s CRM, pipeline, or quote-management records; or (c) was the subject of active business-development activity by Cethos in which Confidential Information about that person or entity was shared with Contractor.</p>
<p>For purposes of this Section, the “end of the relationship” means the later of (a) the last day on which Contractor performed services for Cethos, and (b) the last date on which Cethos paid Contractor for services rendered, in each case as reflected in Cethos’s records.</p>
<p>For purposes of this Section, “solicit” includes any direct or indirect approach, contact, communication, introduction, proposal, marketing, referral, bidding activity, or other conduct intended, or reasonably likely, to obtain business from or establish a direct commercial relationship with a Cethos client or active prospective client outside Cethos.</p>
<p>Contractor shall not, directly or indirectly, circumvent or attempt to circumvent Cethos by seeking to provide services to, contract with, receive work from, invoice, or accept payment from any Cethos client or active prospective client first introduced to Contractor through Cethos, whether in Contractor’s own name or through any affiliate, related party, nominee, intermediary, employee, subcontractor, agent, representative, or other person or entity acting for Contractor’s benefit.</p>
<p>Without limiting the foregoing, Contractor shall not avoid or attempt to avoid the restrictions in this Section by acting through any immediate family member, household member, or business partner, or through any company, partnership, agency, or other entity that (a) Contractor or any immediate family member of Contractor directly or indirectly owns, controls, or manages, or (b) in which Contractor or any such family member holds a material financial or governance interest.</p>
<p>Any business obtained, accepted, performed, or facilitated in breach of this Section by any such person or entity shall be deemed to have been obtained, accepted, performed, or facilitated by Contractor.</p>
<p>For clarity, this Section does not prohibit Contractor from carrying on business generally or engaging in general advertising not directed at Cethos clients or active prospective clients; however, this carve-out does not permit Contractor to solicit, accept, or perform work, directly or indirectly, outside Cethos for any client or active prospective client first introduced to Contractor through Cethos in breach of this Section.</p>
<h3>11. Remedies</h3>
<p>Contractor acknowledges that unauthorized use or disclosure of Confidential Information may cause irreparable harm to Cethos and its clients, and that monetary damages alone may be insufficient. Cethos may seek injunctive or equitable relief in addition to any other remedies available at law or in equity.</p>
<h3>12. Limitation of liability</h3>
<p>Except for liability arising from a party’s wilful misconduct, fraud, or breach of Sections 2, 3, 5, 6, or 10 of this Agreement, each party’s aggregate liability to the other under this Agreement shall not exceed the total fees paid or payable by Cethos to Contractor for the services giving rise to the claim in the twelve (12) months preceding the event giving rise to the claim.</p>
<p>To the maximum extent permitted by applicable law, neither party shall be liable to the other for any loss of profits, loss of business, loss of goodwill, or any indirect, incidental, special, punitive, or consequential damages arising out of or in connection with this Agreement.</p>
<h3>13. Governing law</h3>
<p>This Agreement is governed by the laws of the Province of Alberta, Canada, without regard to its conflict-of-laws principles. The parties attorn to the courts of Alberta for disputes arising out of or relating to this Agreement, unless the parties agree otherwise in writing for a specific client or project.</p>
<h3>14. Electronic acceptance</h3>
<p>By typing a full legal name, checking an acceptance box, applying an electronic signature, or otherwise electronically accepting this Agreement, Contractor confirms that it has read, understood, and agreed to be bound by this Agreement. Such electronic acceptance has the same legal effect as a handwritten signature, to the extent permitted by applicable law.</p>
<h3>15. Independent contractor and authority</h3>
<p>Contractor is an independent contractor and nothing in this Agreement creates any employment, partnership, joint-venture, fiduciary, or agency relationship between Contractor and Cethos. Contractor is solely responsible for all taxes, withholdings, insurance, compensation, and other statutory or contractual obligations of its personnel.</p>
<p>If the person accepting this Agreement does so on behalf of a company, agency, or other entity, that person represents and warrants that they have authority to bind that entity and that references to Contractor include the entity and its personnel engaged on Cethos work.</p>
<h3>16. Entire agreement and precedence</h3>
<p>This Agreement sets out the parties’ confidentiality obligations unless superseded or supplemented by a project-specific agreement, master services agreement, purchase order, or data-processing agreement signed or accepted by the parties. In the event of any conflict, the more specific client, project, or data-protection requirement controls with respect to that project or subject matter.</p>
<h3>17. Severability</h3>
<p>If any provision of this Agreement is held by a court of competent jurisdiction to be invalid, illegal, or unenforceable, that provision shall be deemed modified to the minimum extent necessary to make it valid, legal, and enforceable. If such modification is not possible, that provision shall be severed from this Agreement and the remaining provisions shall remain in full force and effect.</p>
<p>The parties further agree that, with respect to any restrictive covenant in Section 10, if a court determines that any aspect (including duration, scope, or the definition of any term) is unreasonable or overbroad, the court is expressly authorized and requested to modify that aspect to the minimum extent necessary to render the covenant enforceable, and to enforce the covenant as so modified.</p>
<h3>18. Notices</h3>
<p>Any notice required or permitted under this Agreement shall be in writing and given by email to the address designated by the receiving party (and, in the case of Cethos, to legal@cethos.com or such other address as Cethos may designate). Notices are deemed delivered on the date of transmission, provided no automated non-delivery response is received within twenty-four (24) hours. Either party may update its notice email by giving notice in accordance with this Section.</p>
<h3>19. Indemnity</h3>
<p>Contractor shall indemnify and hold harmless Cethos, its Affiliates, Subsidiaries, and their respective directors, officers, employees, and agents (collectively, “Cethos Indemnitees”) from and against any and all claims, losses, damages, liabilities, costs, and expenses (including reasonable legal fees on a full-indemnity basis) suffered or incurred by any Cethos Indemnitee arising out of or in connection with (a) any breach by Contractor of Sections 2, 3, 5, 6, or 10 of this Agreement; (b) the wilful misconduct, fraud, or gross negligence of Contractor or its personnel; or (c) any claim that work product delivered by Contractor infringes the intellectual property rights of a third party, except to the extent caused by Cethos’s own breach or negligence.</p>
<h3>20. Records and audit</h3>
<p>Contractor shall maintain reasonable records of its security practices, subcontractor confidentiality undertakings, and AI and machine-translation tools used in the performance of services, for a period of not less than two (2) years following completion of the relevant services. On reasonable prior written notice and not more than once in any twelve-month period (except in response to a suspected breach or a client or regulator request), Cethos may request, and Contractor shall provide, written information sufficient to verify Contractor’s compliance with Sections 3, 4, 5, and 6. Cethos shall conduct any such review in a manner that minimizes disruption to Contractor’s business.</p>
<h3>21. Miscellaneous</h3>
<p>(a) Waiver. No failure or delay by either party in exercising any right under this Agreement constitutes a waiver of that right, and no single or partial exercise of any right precludes any other or further exercise.</p>
<p>(b) Assignment. Contractor shall not assign or transfer this Agreement, in whole or in part, without Cethos’s prior written consent. Cethos may assign this Agreement to an Affiliate or Subsidiary, or to a successor in connection with a merger, reorganization, or sale of all or substantially all of its assets, without consent.</p>
<p>(c) Further assurances. Each party shall execute such further documents and take such further actions as may reasonably be required to give effect to this Agreement.</p>
<p>(d) Headings. Headings are for convenience only and do not affect interpretation.</p>
<p>(e) Counterparts. This Agreement may be accepted in counterparts, including by electronic means, each of which is an original and all of which together constitute one and the same instrument.</p>""",
    },
]


def html_unescape(s: str) -> str:
    """Cheap entity decode — we only emit &amp; &lt; &gt; &quot; in templates."""
    return (
        s.replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
        .replace("&#39;", "'")
    )


class ParaCollector(HTMLParser):
    """Walks the narrow HTML schema and yields (tag, [(bold, text), ...]).

    Bold runs are produced for <strong>/<b>; everything else is plain.
    """

    def __init__(self) -> None:
        super().__init__()
        self.paragraphs: list[tuple[str, list[tuple[bool, str]]]] = []
        self._current_tag: str | None = None
        self._current_runs: list[tuple[bool, str]] = []
        self._bold_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in ("h2", "h3", "p", "li"):
            self._flush()
            self._current_tag = tag
            self._current_runs = []
        elif tag in ("strong", "b"):
            self._bold_depth += 1

    def handle_endtag(self, tag):
        if tag in ("h2", "h3", "p", "li"):
            self._flush()
        elif tag in ("strong", "b"):
            self._bold_depth = max(0, self._bold_depth - 1)

    def handle_data(self, data):
        if self._current_tag is None:
            return
        text = html_unescape(data)
        if not text.strip() and not self._current_runs:
            return
        self._current_runs.append((self._bold_depth > 0, text))

    def _flush(self):
        if self._current_tag is None:
            return
        # Skip empty paragraphs.
        if any(run[1].strip() for run in self._current_runs):
            self.paragraphs.append((self._current_tag, self._current_runs))
        self._current_tag = None
        self._current_runs = []

    def close(self):
        self._flush()
        super().close()


def build_docx(template: dict, out_path: Path) -> None:
    doc = Document()

    # Page margins — slightly tighter than Word's 1in default.
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    # Default style: 11pt Times for body, line spacing 1.15.
    base = doc.styles["Normal"]
    base.font.name = "Times New Roman"
    base.font.size = Pt(11)

    # Title block.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CETHOS SOLUTIONS INC.")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(template["title"])
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Version {template['version_label']}  ·  Effective {template['effective_from']}"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()  # spacing

    # Parse body.
    parser = ParaCollector()
    parser.feed(template["body_html"])
    parser.close()

    for tag, runs in parser.paragraphs:
        # First h2 in the doc is the agreement title — we already
        # rendered it as the page subtitle above, skip the duplicate.
        if tag == "h2" and not any(
            p.text and p.text.strip() not in (
                "CETHOS SOLUTIONS INC.",
                template["title"],
                meta.text.strip(),
            )
            for p in doc.paragraphs
        ):
            continue

        para = doc.add_paragraph()
        if tag == "h2":
            for bold, text in runs:
                r = para.add_run(text)
                r.font.name = "Calibri"
                r.font.size = Pt(13)
                r.bold = True
            continue
        if tag == "h3":
            for bold, text in runs:
                r = para.add_run(text)
                r.font.name = "Calibri"
                r.font.size = Pt(11.5)
                r.bold = True
                r.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
            continue
        if tag == "li":
            para.style = doc.styles["List Bullet"]
            for bold, text in runs:
                r = para.add_run(text)
                r.bold = bold
            continue
        # <p>
        for bold, text in runs:
            r = para.add_run(text)
            r.bold = bold

    # Acceptance footer note — this template is executed electronically
    # via the vendor portal; the .docx is a reference copy of the text
    # that appears on screen at sign time.
    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run(
        "This document is the reference copy of the Cethos Translator NDA. "
        "Vendors execute it electronically via the Cethos vendor portal "
        "(https://vendor.cethos.com/nda); the signed copy includes a "
        "verification audit log and is downloadable as a PDF."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.save(out_path)
    print(f"wrote {out_path}")


def main() -> int:
    repo_root = Path(__file__).resolve().parents[1]
    out_dir = repo_root / "docs" / "nda-templates"
    out_dir.mkdir(parents=True, exist_ok=True)
    for template in TEMPLATES:
        out = out_dir / f"cethos-nda-{template['version_label']}.docx"
        build_docx(template, out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
